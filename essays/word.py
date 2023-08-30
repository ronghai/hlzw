#! /usr/bin/env python3
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm
from docx.shared import  Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import  qn
from docx.shared import Inches
from docx.shared import Cm

def a4():
    document = Document()
    section = document.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.top_margin = Mm(25.4)
    section.bottom_margin = Mm(25.4)
    section.left_margin = Mm(31.7)
    section.right_margin = Mm(31.7) 
    section.header_distance = Mm(15.0)
    section.footer_distance = Mm(17.5)
    return document

#WORD中字号、磅值，两者具体的对应关系大约如下吧：
#字号‘八号’对应磅值5
#字号‘七号’对应磅值5.5
#字号‘小六’对应磅值6.5
#字号‘六号’对应磅值7.5
#字号‘小五’对应磅值9
#字号‘五号’对应磅值10.5
#字号‘小四’对应磅值12
#字号‘四号’对应磅值14
#字号‘小三’对应磅值15
#字号‘三号’对应磅值16
#字号‘小二’对应磅值18
#字号‘二号’对应磅值22
#字号‘小一’对应磅值24
#字号‘一号’对应磅值26
#字号‘小初’对应磅值36
#字号‘初号’对应磅值42


def zuowen(document, no, name, subject, body):
    #document = a4()
    document.styles['Normal'].font.name = u'宋体' #设置西文字体
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体') #设置中文字体使用字体2->宋体
    document.styles['Normal'].font.size=Pt(14)

    head = document.add_paragraph()
    head.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    head.paragraph_format.space_before = Pt(0)
    head.paragraph_format.space_after = Pt(0)
    head.paragraph_format.line_spacing = 1 #Pt(16)
    run = head.add_run(subject)
    run.font.size=Pt(16) #字号‘三号’对应磅值16
    run.bold = True
    run.font.name = u'宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    
    paragraph = document.add_paragraph("")
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1 #Pt(14)
    run = paragraph.add_run("%d号 %s" % (no, name))
    run.font.size=Pt(14) #字号‘四号’对应磅值14
    run.bold = True
    run.font.name = u'楷体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'楷体') 

    paragraphes = body.splitlines()
    for para in paragraphes:
        p = para.strip()
        if len(p) == 0:
            continue
        paragraph = document.add_paragraph(para)
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = Pt(28)
        paragraph.paragraph_format.first_line_indent = Pt(14) * 2
    return document

if __name__ == "__main__":
    body = '''
我有一个非常要好的朋友，他的个子非常高，有一头乌黑的头发，一双浓密的眉毛下面有一双亮晶晶的眼睛，一张小小的嘴巴特别能说会道。一身小麦色的皮肤，看起来很健康。
他是一个很爱笑的人，有一次，我说了个米小圈的笑话，他笑得肚子都疼了；有时候老师表扬他，他就会嘻嘻地笑；有时他会一个人偷偷地傻笑。正因为他爱笑，给人的感觉他很自信，每次考试，他都很冷静。考试结果都很好。
考试考得好是因为他平时作业及时完成，认真听老师说话。老师都夸奖他学习认真，学习效率高。每当他遇到困难，他总会请教别人，寻求帮助。勤学好问是他优秀的品质。除了学习之外，他还乐于帮助别人，我有一次橡皮不见了，他就把橡皮借给了我。
他还有很多爱好，钢琴他已经是六级了，他还是我们班的篮球能手，他跑步每次都是第五六名，这就是我的好朋友。
我要和他做一辈子最好的朋友。
'''
    document = zuowen(25, "李晟睿", "猜猜他是谁", body)
    document.save('25号 李晟睿.2.docx')
